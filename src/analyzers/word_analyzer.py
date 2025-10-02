"""
Word document analyzer for .doc and .docx files
"""

import logging
from pathlib import Path
from typing import Dict, Any

try:
    from docx import Document
    import docx.oxml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

from . import BaseAnalyzer

class WordAnalyzer(BaseAnalyzer):
    """Analyzer for Word documents (.docx and .doc files)"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx not available. Advanced .docx analysis will be limited.")
        if not DOCX2TXT_AVAILABLE:
            self.logger.warning("docx2txt not available. .doc file support will be limited.")
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text content from Word documents"""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.doc':
            return self._extract_text_from_doc(file_path)
        elif file_extension == '.docx':
            return self._extract_text_from_docx(file_path)
        else:
            return f"Unsupported file extension: {file_extension}"
    
    def _extract_text_from_doc(self, file_path: Path) -> str:
        """Extract text from legacy .doc files with fallback to text extraction"""
        
        # First, try to detect if this is actually a text/HTML file with .doc extension
        try:
            with open(file_path, 'rb') as f:
                first_bytes = f.read(16)
            
            # Check if it's not a real .doc file (should start with D0CF11E0 for OLE format)
            if not first_bytes.startswith(b'\xd0\xcf\x11\xe0'):
                # Likely a text/HTML file with .doc extension (common with Confluence exports)
                return self._extract_text_as_plaintext(file_path)
                
        except Exception as e:
            self.logger.warning(f"Could not check file format for {file_path}: {e}")
        
        # Try docx2txt for actual .doc files
        if not DOCX2TXT_AVAILABLE:
            return "Legacy .doc files require docx2txt package. Please install with: pip install docx2txt"
        
        try:
            # Use docx2txt for .doc files
            text_content = docx2txt.process(str(file_path))
            return text_content if text_content else "No text content found in .doc file"
            
        except Exception as e:
            self.logger.warning(f"docx2txt failed for {file_path}: {e}")
            # Fallback to text extraction
            return self._extract_text_as_plaintext(file_path)
    
    def _extract_text_as_plaintext(self, file_path: Path) -> str:
        """Fallback method to extract text content treating file as plain text"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Clean up HTML if present (common in Confluence exports)
            if '<html' in content.lower() or '<!doctype' in content.lower():
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(content, 'html.parser')
                    text_content = soup.get_text(separator='\n', strip=True)
                    return text_content if text_content else content
                except ImportError:
                    # BeautifulSoup not available, extract basic HTML
                    import re
                    # Remove HTML tags
                    clean_text = re.sub(r'<[^>]+>', ' ', content)
                    # Clean up whitespace
                    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                    return clean_text
            
            return content
            
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                return f"Error reading file as text: {str(e)}"
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from .docx files with fallback options"""
        # Try python-docx first for better formatting
        if DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                text_content = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                result = "\n".join(text_content)
                if result:
                    return result
                    
            except Exception as e:
                self.logger.warning(f"python-docx failed for {file_path}, trying fallback: {e}")
        
        # Fallback to docx2txt for .docx files
        if DOCX2TXT_AVAILABLE:
            try:
                text_content = docx2txt.process(str(file_path))
                return text_content if text_content else "No text content found in .docx file"
            except Exception as e:
                self.logger.error(f"docx2txt also failed for {file_path}: {e}")
                return f"Error reading .docx file: {str(e)}"
        
        return "Both python-docx and docx2txt are unavailable. Cannot process .docx files."
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Word documents"""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.doc':
            return self._extract_metadata_from_doc(file_path)
        elif file_extension == '.docx':
            return self._extract_metadata_from_docx(file_path)
        else:
            return {'error': f'Unsupported file extension: {file_extension}'}
    
    def _extract_metadata_from_doc(self, file_path: Path) -> Dict[str, Any]:
        """Extract basic metadata from .doc files"""
        metadata = {
            'file_type': 'word_legacy_doc',
            'supports_full_metadata': False,
            'extraction_method': 'auto_detect'
        }
        
        try:
            # Check file format
            with open(file_path, 'rb') as f:
                first_bytes = f.read(16)
            
            if not first_bytes.startswith(b'\xd0\xcf\x11\xe0'):
                metadata['actual_format'] = 'text_html_confluence_export'
                metadata['extraction_method'] = 'text_fallback'
            else:
                metadata['actual_format'] = 'microsoft_ole_document'
                metadata['extraction_method'] = 'docx2txt'
        except Exception:
            metadata['actual_format'] = 'unknown'
        
        try:
            # Basic file stats
            stat = file_path.stat()
            metadata.update({
                'file_size': stat.st_size,
                'created': str(stat.st_ctime),
                'modified': str(stat.st_mtime),
            })
            
            # Try to extract text to get word count
            if metadata.get('actual_format') == 'text_html_confluence_export':
                # Use text extraction method
                text_content = self._extract_text_as_plaintext(file_path)
            else:
                # Use docx2txt method
                if DOCX2TXT_AVAILABLE:
                    text_content = docx2txt.process(str(file_path))
                else:
                    text_content = self._extract_text_as_plaintext(file_path)
            
            if text_content:
                word_count = len(text_content.split())
                char_count = len(text_content)
                metadata.update({
                    'word_count': word_count,
                    'character_count': char_count,
                    'has_content': True
                })
            else:
                metadata.update({
                    'word_count': 0,
                    'character_count': 0,
                    'has_content': False
                })
                
        except Exception as e:
            self.logger.error(f"Error extracting .doc metadata from {file_path}: {e}")
            metadata['error'] = str(e)
            
        return metadata
    
    def _extract_metadata_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract comprehensive metadata from .docx files"""
        if not DOCX_AVAILABLE:
            # Fallback to basic metadata
            return self._extract_basic_docx_metadata(file_path)
        
        try:
            doc = Document(file_path)
            
            metadata = {
                'file_type': 'word_docx',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'extraction_method': 'python-docx'
            }
            
            # Extract core properties if available
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata['title'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                metadata['subject'] = doc.core_properties.subject
            if hasattr(doc.core_properties, 'keywords') and doc.core_properties.keywords:
                metadata['keywords'] = doc.core_properties.keywords
            if hasattr(doc.core_properties, 'comments') and doc.core_properties.comments:
                metadata['comments'] = doc.core_properties.comments
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                metadata['created'] = str(doc.core_properties.created)
            if hasattr(doc.core_properties, 'modified') and doc.core_properties.modified:
                metadata['modified'] = str(doc.core_properties.modified)
            if hasattr(doc.core_properties, 'last_modified_by') and doc.core_properties.last_modified_by:
                metadata['last_modified_by'] = doc.core_properties.last_modified_by
            
            # Count different elements
            styles_used = set()
            for paragraph in doc.paragraphs:
                if paragraph.style:
                    styles_used.add(paragraph.style.name)
            
            metadata['styles_used'] = list(styles_used)
            metadata['style_count'] = len(styles_used)
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting .docx metadata from {file_path}: {e}")
            # Try fallback method
            return self._extract_basic_docx_metadata(file_path)
    
    def _extract_basic_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Fallback metadata extraction for .docx files using docx2txt"""
        metadata = {
            'file_type': 'word_docx',
            'extraction_method': 'docx2txt_fallback'
        }
        
        if not DOCX2TXT_AVAILABLE:
            metadata['error'] = 'No extraction method available'
            return metadata
        
        try:
            # Basic file stats
            stat = file_path.stat()
            metadata.update({
                'file_size': stat.st_size,
                'created': str(stat.st_ctime),
                'modified': str(stat.st_mtime),
            })
            
            # Try to extract text to get word count
            text_content = docx2txt.process(str(file_path))
            if text_content:
                word_count = len(text_content.split())
                char_count = len(text_content)
                line_count = len(text_content.split('\n'))
                metadata.update({
                    'word_count': word_count,
                    'character_count': char_count,
                    'line_count': line_count,
                    'has_content': True
                })
            else:
                metadata.update({
                    'word_count': 0,
                    'character_count': 0,
                    'line_count': 0,
                    'has_content': False
                })
                
        except Exception as e:
            self.logger.error(f"Error extracting basic .docx metadata from {file_path}: {e}")
            metadata['error'] = str(e)
            
        return metadata